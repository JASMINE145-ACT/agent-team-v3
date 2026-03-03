# -*- coding: utf-8 -*-
"""Extract text from Coze docx for conversion."""
import os
import sys

base = os.path.dirname(os.path.abspath(__file__))
path = os.path.join(base, "Coze助力PT. VANTSING企业微信+ACCURATE系统全业务流程实施方案.docx")
out_path = os.path.join(base, "Coze企划提取.txt")
if not os.path.exists(path):
    print("File not found:", path, file=sys.stderr)
    sys.exit(1)
from docx import Document
doc = Document(path)
lines = []
for p in doc.paragraphs:
    lines.append(p.text)
for t in doc.tables:
    for row in t.rows:
        lines.append(" | ".join(c.text.strip() for c in row.cells))
with open(out_path, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))
print("Written to", out_path)
